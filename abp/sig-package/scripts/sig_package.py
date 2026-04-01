#!/usr/bin/env python3
"""
Signature Package Extractor

Extracts signature pages from a folder of closing documents (.docx and .pdf)
and assembles them into bookmarked PDF signature packages.

Usage:
    python sig_package.py <input_folder> [--auto-run] [--project-name "Name"]

Dependencies:
    pip install python-docx pypdf thefuzz openpyxl Pillow --break-system-packages
    pandoc must be installed for .docx-to-PDF conversion (preferred).
    LibreOffice is a fallback if pandoc is unavailable.
"""

import argparse
import json
import os
import re
import subprocess
import sys
import tempfile
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

try:
    from docx import Document as DocxDocument
    from thefuzz import fuzz
    from pypdf import PdfReader, PdfWriter
    import openpyxl
except ImportError as e:
    print(f"Missing dependency: {e}")
    print("Install with: pip install python-docx pypdf thefuzz openpyxl Pillow --break-system-packages")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class ChecklistEntry:
    """A single document entry from the closing checklist."""
    order: int
    doc_name: str
    signatories: list[str] = field(default_factory=list)
    matched_file: Optional[str] = None
    match_score: int = 0


@dataclass
class SigPageCandidate:
    """A candidate signature page detected in a document."""
    source_file: str
    page_number: int  # 0-indexed
    markers_found: list[str] = field(default_factory=list)
    marker_score: float = 0.0
    structural_score: float = 0.0
    total_score: float = 0.0
    is_notary: bool = False
    signatory_names: list[str] = field(default_factory=list)
    associated_signatory: Optional[str] = None
    detection_layer: str = ""  # "footer", "keyword", "structural"


@dataclass
class SignatoryInfo:
    """Normalized signatory information."""
    canonical_name: str
    name_variants: list[str] = field(default_factory=list)
    pages: list[SigPageCandidate] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Phase 1: Parse the Closing Checklist
# ---------------------------------------------------------------------------

CHECKLIST_PATTERNS = [
    "checklist", "closing list", "doc list", "document list",
    "closing schedule", "closing index"
]


def extract_project_name_from_filenames(folder: Path) -> Optional[str]:
    """Extract the project name from the common filename prefix.

    Files follow the convention: [Project Name] - [Document Name].docx
    E.g., '307-313 W Martin - Deed of Trust.docx'
    The project name is everything before the first ' - ' separator,
    derived from the most common prefix across all .docx/.pdf files.
    """
    doc_files = [
        f for f in folder.iterdir()
        if f.suffix.lower() in (".docx", ".pdf") and f.is_file()
        and not f.name.startswith("~$")
        and "SigPackage" not in f.name
    ]
    if not doc_files:
        return None

    # Extract prefixes (everything before first " - ")
    prefixes = []
    for f in doc_files:
        parts = f.stem.split(" - ", 1)
        if len(parts) == 2:
            prefixes.append(parts[0].strip())

    if not prefixes:
        return None

    # Find the most common prefix
    from collections import Counter
    counts = Counter(prefixes)
    most_common = counts.most_common(1)[0]

    # Only use it if a majority of files share this prefix
    if most_common[1] >= len(doc_files) * 0.5:
        return most_common[0]

    return None


def find_checklist_file(folder: Path) -> Optional[Path]:
    """Locate the closing checklist file in the folder."""
    candidates = []
    for f in folder.iterdir():
        if f.suffix.lower() in (".docx", ".xlsx", ".xls"):
            name_lower = f.stem.lower()
            for pattern in CHECKLIST_PATTERNS:
                if pattern in name_lower:
                    candidates.append((f, pattern))
                    break

    if not candidates:
        return None
    # Prefer "checklist" match
    for f, pattern in candidates:
        if pattern == "checklist":
            return f
    return candidates[0][0]


def parse_checklist_docx(filepath: Path) -> tuple[str, list[ChecklistEntry]]:
    """Parse a Word document closing checklist with tables."""
    doc = DocxDocument(str(filepath))
    project_name = ""
    entries = []

    # Try to extract project name from paragraphs before tables
    # Strategy: look for explicit "Project:" or "Loan:" labels first, then fall back
    # to title-like text. Skip generic titles like "Signature Checklist" or "Closing Checklist".
    generic_titles = {"signature checklist", "closing checklist", "document list",
                      "closing list", "doc list", "checklist"}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Look for explicit project/loan label (e.g., "Project: Martin Refinance Loan")
        label_match = re.match(
            r'(?:project|loan|transaction|deal|property|re)\s*[:]\s*(.+)',
            text, re.IGNORECASE
        )
        if label_match:
            project_name = label_match.group(1).strip().strip('"').strip("'")
            break

        # Look for quoted project name in text (e.g., the "Martin Refinance Loan")
        quoted_match = re.search(r'["\u201c]([^"\u201d]+)["\u201d]\s*(?:loan|project|transaction)?', text, re.IGNORECASE)
        if quoted_match and quoted_match.group(1).lower() not in generic_titles:
            project_name = quoted_match.group(1).strip()
            break

    # If no explicit label found, try title-like paragraphs (but skip generic titles)
    if not project_name:
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading") or (para.runs and any(r.bold for r in para.runs)):
                cleaned = re.sub(
                    r'\s*[-–—]\s*(closing\s+)?(checklist|doc(ument)?\s+list|schedule|index).*$',
                    '', text, flags=re.IGNORECASE
                ).strip()
                if cleaned and cleaned.lower() not in generic_titles:
                    project_name = cleaned
                    break

    # Parse tables
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        # Identify header row
        header_row = table.rows[0]
        headers = [cell.text.strip().lower() for cell in header_row.cells]

        # Find document name column
        doc_col = None
        sig_cols = []
        for i, h in enumerate(headers):
            if any(kw in h for kw in ["document", "description", "instrument", "name", "item", "title"]):
                if doc_col is None:
                    doc_col = i
            if any(kw in h for kw in ["signatory", "signer", "signature", "execution",
                                       "signed by", "parties", "borrower", "lender",
                                       "guarantor", "party"]):
                sig_cols.append(i)

        if doc_col is None:
            # Try first text-heavy column
            doc_col = 0 if len(headers) > 0 else None

        if doc_col is None:
            continue

        # Parse data rows
        order = 0
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if doc_col >= len(cells):
                continue

            doc_name = cells[doc_col]
            if not doc_name:
                continue

            # Skip exhibit/schedule sub-headers
            if re.match(r'^(exhibit|schedule|attachment|annex)\b', doc_name, re.IGNORECASE):
                continue

            # Clean document name
            doc_name = re.sub(r'^[\d]+[.\)]\s*', '', doc_name).strip()
            doc_name = re.sub(r'^[A-Z][.\)]\s*', '', doc_name).strip()

            # Extract signatory names
            sigs = []
            for sc in sig_cols:
                if sc < len(cells):
                    cell_text = cells[sc]
                    # Skip checkmarks and simple markers
                    if cell_text.lower() in ("x", "yes", "✓", "✔", ""):
                        # This is a role-based column — use header as role hint, not as name
                        continue
                    # Split on common delimiters
                    for name in re.split(r'[;,\n]', cell_text):
                        name = name.strip()
                        if name and len(name) > 1:
                            sigs.append(name)

            order += 1
            entries.append(ChecklistEntry(
                order=order,
                doc_name=doc_name,
                signatories=sigs
            ))

    # Fallback project name from filename
    if not project_name:
        project_name = re.sub(
            r'[-_]?(closing[-_]?)?(checklist|doc[-_]?list|schedule|index).*$',
            '', filepath.stem, flags=re.IGNORECASE
        ).strip().replace('-', ' ').replace('_', ' ')

    return project_name, entries


def parse_checklist_xlsx(filepath: Path) -> tuple[str, list[ChecklistEntry]]:
    """Parse an Excel closing checklist."""
    wb = openpyxl.load_workbook(str(filepath), data_only=True)
    ws = wb.active
    project_name = ""
    entries = []

    # Find header row and project name
    header_row_idx = None
    for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=10, values_only=False), start=1):
        cell_values = [str(cell.value).strip().lower() if cell.value else "" for cell in row]

        # Check if this looks like a header row
        doc_indicators = ["document", "description", "instrument", "name", "item", "title"]
        if any(any(kw in cv for kw in doc_indicators) for cv in cell_values if cv):
            header_row_idx = row_idx
            break

        # Otherwise might be project name
        for cell in row:
            if cell.value and not project_name:
                text = str(cell.value).strip()
                if len(text) > 3 and len(text) < 150:
                    cleaned = re.sub(
                        r'\s*[-–—]\s*(closing\s+)?(checklist|doc(ument)?\s+list).*$',
                        '', text, flags=re.IGNORECASE
                    ).strip()
                    if cleaned:
                        project_name = cleaned

    if header_row_idx is None:
        header_row_idx = 1  # Assume first row is header

    # Parse headers
    headers = []
    for cell in ws[header_row_idx]:
        headers.append(str(cell.value).strip().lower() if cell.value else "")

    # Find columns
    doc_col = None
    sig_cols = []
    for i, h in enumerate(headers):
        if any(kw in h for kw in ["document", "description", "instrument", "name", "item", "title"]):
            if doc_col is None:
                doc_col = i
        if any(kw in h for kw in ["signatory", "signer", "signature", "execution",
                                   "signed by", "parties", "borrower", "lender",
                                   "guarantor", "party"]):
            sig_cols.append(i)

    if doc_col is None:
        doc_col = 0

    # Parse data rows
    order = 0
    for row in ws.iter_rows(min_row=header_row_idx + 1, values_only=True):
        if not row or all(cell is None for cell in row):
            break

        doc_name = str(row[doc_col]).strip() if doc_col < len(row) and row[doc_col] else ""
        if not doc_name or doc_name == "None":
            continue

        # Skip exhibit/schedule sub-headers
        if re.match(r'^(exhibit|schedule|attachment|annex)\b', doc_name, re.IGNORECASE):
            continue

        # Clean document name
        doc_name = re.sub(r'^[\d]+[.\)]\s*', '', doc_name).strip()

        # Extract signatory names
        sigs = []
        for sc in sig_cols:
            if sc < len(row) and row[sc]:
                cell_text = str(row[sc]).strip()
                if cell_text.lower() in ("x", "yes", "✓", "✔", "none", ""):
                    continue
                for name in re.split(r'[;,\n]', cell_text):
                    name = name.strip()
                    if name and len(name) > 1:
                        sigs.append(name)

        order += 1
        entries.append(ChecklistEntry(
            order=order,
            doc_name=doc_name,
            signatories=sigs
        ))

    # Fallback project name
    if not project_name:
        project_name = re.sub(
            r'[-_]?(closing[-_]?)?(checklist|doc[-_]?list|schedule|index).*$',
            '', filepath.stem, flags=re.IGNORECASE
        ).strip().replace('-', ' ').replace('_', ' ')

    return project_name, entries


def parse_checklist(filepath: Path) -> tuple[str, list[ChecklistEntry]]:
    """Parse the closing checklist regardless of format."""
    if filepath.suffix.lower() == ".docx":
        return parse_checklist_docx(filepath)
    elif filepath.suffix.lower() in (".xlsx", ".xls"):
        return parse_checklist_xlsx(filepath)
    else:
        raise ValueError(f"Unsupported checklist format: {filepath.suffix}")


def fuzzy_match_files(entries: list[ChecklistEntry], folder: Path) -> list[ChecklistEntry]:
    """Fuzzy-match checklist entries to files in the folder."""
    # Gather all candidate files
    candidate_files = [
        f for f in folder.iterdir()
        if f.suffix.lower() in (".docx", ".pdf") and f.is_file()
    ]

    def clean_name(name: str) -> str:
        """Normalize a name for fuzzy matching."""
        name = re.sub(r'\.(docx|pdf|xlsx?)$', '', name, flags=re.IGNORECASE)
        name = re.sub(r'^[\d]+[-_.\s]+', '', name)  # Leading numbers
        # Strip common project-name prefixes like "307-313 W Martin - "
        # Pattern: optional address/numbers, then words, then a separator (- or –)
        name = re.sub(r'^[\d\-]+\s+[\w\s]+?\s*[-–—]\s*', '', name)
        name = re.sub(r'\s*[-–—]\s*(draft|final|execution\s*copy|v\d+|clean).*$', '',
                       name, flags=re.IGNORECASE)
        name = re.sub(r'\s*\(v?\d+\)\s*$', '', name, flags=re.IGNORECASE)
        name = name.replace('_', ' ').replace('-', ' ')
        return name.strip().lower()

    used_files = set()
    for entry in entries:
        best_score = 0
        best_file = None
        cleaned_entry = clean_name(entry.doc_name)

        for f in candidate_files:
            if f.name in used_files:
                continue
            # Skip checklist file itself
            if any(p in f.stem.lower() for p in CHECKLIST_PATTERNS):
                continue
            cleaned_file = clean_name(f.stem)
            score = fuzz.token_sort_ratio(cleaned_entry, cleaned_file)
            if score > best_score:
                best_score = score
                best_file = f

        if best_file and best_score >= 70:
            entry.matched_file = str(best_file)
            entry.match_score = best_score
            used_files.add(best_file.name)

    return entries


# ---------------------------------------------------------------------------
# Phase 2: Signature Page Detection
# ---------------------------------------------------------------------------

# Primary markers
PRIMARY_MARKERS = [
    (r"in\s+witness\s+whereof", "IN WITNESS WHEREOF"),
    (r"executed\s+(as\s+of|effective|this)", "EXECUTED as of/effective/this"),
    (r"\[seal\]", "[SEAL]"),
    (r"witness\s+my\s+hand", "WITNESS my hand"),
    (r"signed[,\s]+(sealed|and\s+sealed)", "SIGNED, SEALED"),
    (r"(?:^|\s)witnesses?\s*:", "WITNESS(ES):"),
    (r"\[?\s*signature\s+page\s+follows?\s*\]?", "[Signature Page Follows]"),
]

# Signature block markers
SIG_BLOCK_MARKERS = [
    (r"(?:^|\n)\s*by:\s*[_\s\n]", "By: ___"),
    (r"(?:^|\n)\s*by:\s*\w", "By: [name]"),
    (r"(?:^|\n)\s*(?:print(?:ed)?\s+)?name:\s*[_\s\n]", "Name: ___"),
    (r"(?:^|\n)\s*(?:print(?:ed)?\s+)?name:\s*\w", "Name: [filled]"),
    (r"(?:^|\n)\s*printed:\s*[_\s\n]", "Printed: ___"),
    (r"(?:^|\n)\s*printed:\s*\w", "Printed: [filled]"),
    (r"(?:^|\n)\s*its:\s*[_\s\n]", "Its: ___"),
    (r"(?:^|\n)\s*its:\s*\w", "Its: [filled]"),
    (r"_{5,}", "Underscores (5+)"),
    (r"authorized\s+(?:signatory|representative|officer)", "Authorized Signatory/Rep"),
    (r"(?:^|\n)\s*signature\s*:", "Signature:"),
]

# Notary markers
NOTARY_MARKERS = [
    (r"state\s+of", "STATE OF"),
    (r"county\s+of", "COUNTY OF"),
    (r"personally\s+appeared", "personally appeared"),
    (r"notary\s+public", "Notary Public"),
    (r"sworn\s+and\s+subscribed", "sworn and subscribed"),
    (r"acknowledged\s+before\s+me", "acknowledged before me"),
    (r"my\s+commission\s+expires", "My commission expires"),
    (r"(?:^|\s)(?:acknowledgment|jurat)\s*$", "ACKNOWLEDGMENT/JURAT"),
    (r"subscribed\s+and\s+sworn", "subscribed and sworn"),
    (r"before\s+me,?\s+the\s+undersigned", "before me, the undersigned"),
]

# Footer patterns (Layer 0)
FOOTER_PATTERNS = [
    r"signature\s*[-–—]\s*.+",  # "Signature - Doc Name"
    r"signature\s+page",
    r"sig\s+page",
    r"execution\s+page",
    r"^signature$",
]


def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Optional[Path]:
    """Convert a .docx file to PDF using LibreOffice via the soffice wrapper.

    Uses the soffice_wrapper.py shim that handles sandboxed environments where
    AF_UNIX sockets are blocked. This produces high-fidelity PDFs that preserve
    the original Word formatting, unlike pandoc which completely reflows the document.
    """
    pdf_name = docx_path.stem + ".pdf"
    pdf_path = output_dir / pdf_name

    # Use the soffice wrapper (handles sandbox socket restrictions via LD_PRELOAD shim)
    wrapper_path = Path(__file__).parent / "soffice_wrapper.py"
    try:
        result = subprocess.run(
            ["python3", str(wrapper_path),
             "--headless", "--norestore", "--convert-to", "pdf",
             "--outdir", str(output_dir), str(docx_path)],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and pdf_path.exists():
            return pdf_path
        print(f"  Warning: LibreOffice conversion failed for {docx_path.name}: {result.stderr}")
    except subprocess.TimeoutExpired:
        print(f"  Warning: LibreOffice conversion timed out for {docx_path.name}")
    except FileNotFoundError:
        print("  Error: soffice_wrapper.py not found. Ensure it is in the scripts/ directory.")

    # Fallback: try direct LibreOffice (may work outside sandbox)
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(output_dir), str(docx_path)],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode == 0 and pdf_path.exists():
            return pdf_path
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass

    print(f"  Error: Could not convert {docx_path.name} to PDF. LibreOffice is required.")
    return None


def extract_page_text(pdf_reader: PdfReader, page_idx: int) -> str:
    """Extract text from a single PDF page."""
    try:
        return pdf_reader.pages[page_idx].extract_text() or ""
    except Exception:
        return ""


def get_footer_text(page_text: str) -> str:
    """Extract the bottom portion of page text (approximate footer area)."""
    lines = page_text.strip().split('\n')
    if len(lines) <= 3:
        return page_text
    # Take bottom ~10% of lines, minimum 2 lines
    footer_lines = max(2, len(lines) // 10)
    return '\n'.join(lines[-footer_lines:])


def check_footer_match(footer_text: str) -> Optional[str]:
    """Check if footer matches signature page patterns (Layer 0)."""
    footer_lower = footer_text.lower().strip()
    for pattern in FOOTER_PATTERNS:
        if re.search(pattern, footer_lower, re.MULTILINE | re.IGNORECASE):
            return f"Footer: {footer_text.strip()[-50:]}"
    return None


def score_page(page_text: str, page_idx: int, total_pages: int,
               avg_text_length: float, prev_is_candidate: bool) -> SigPageCandidate:
    """Score a single page for signature page likelihood."""
    candidate = SigPageCandidate(source_file="", page_number=page_idx)
    text_lower = page_text.lower()

    # Layer 0: Footer check
    footer_text = get_footer_text(page_text)
    footer_match = check_footer_match(footer_text)
    if footer_match:
        candidate.markers_found.append(footer_match)
        candidate.detection_layer = "footer"
        candidate.total_score = 100  # Automatic qualify
        return candidate

    # Layer 1: Keyword detection
    primary_score = 0
    for pattern, label in PRIMARY_MARKERS:
        if re.search(pattern, text_lower):
            candidate.markers_found.append(label)
            primary_score += 2

    sig_block_score = 0
    for pattern, label in SIG_BLOCK_MARKERS:
        matches = re.findall(pattern, text_lower, re.MULTILINE)
        if matches:
            candidate.markers_found.append(f"{label} (×{len(matches)})")
            sig_block_score += len(matches)  # Count each occurrence

    notary_score = 0
    is_notary = False
    # Require STATE OF / COUNTY OF in notary-block format (start of line, or
    # followed by ")", "§", "ss").  Mid-sentence occurrences like "the laws
    # of the State of New York" are NOT notary blocks.
    notary_state_re = r'(?:^|(?<=\n))\s*state\s+of\b|state\s+of\s*[\)§]'
    notary_county_re = r'(?:^|(?<=\n))\s*county\s+of\b|county\s+of\s*[\)§]'
    has_state_of = bool(re.search(notary_state_re, text_lower, re.MULTILINE))
    has_county_of = bool(re.search(notary_county_re, text_lower, re.MULTILINE))
    if has_state_of and has_county_of:
        is_notary = True
        notary_score += 2
        candidate.markers_found.append("STATE OF + COUNTY OF")

    for pattern, label in NOTARY_MARKERS[2:]:  # Skip STATE OF and COUNTY OF (already checked)
        if re.search(pattern, text_lower, re.MULTILINE):
            candidate.markers_found.append(label)
            notary_score += 1

    if not is_notary and notary_score >= 2:
        has_notary_public = bool(re.search(r"notary\s+public", text_lower))
        if has_notary_public:
            is_notary = True

    candidate.is_notary = is_notary
    candidate.marker_score = primary_score + sig_block_score

    # Layer 2: Structural analysis (only if at least 1 keyword point)
    structural = 0
    if primary_score > 0 or sig_block_score > 0 or notary_score > 0:
        # Short document adjustment
        if total_pages <= 5:
            if page_idx >= total_pages - 2:
                structural += 1
        else:
            if page_idx >= total_pages * 0.8:
                structural += 2
            if page_idx >= total_pages - 3:
                structural += 1

        if len(page_text) < avg_text_length * 0.7:
            structural += 1

        by_count = len(re.findall(r"(?:^|\s)by:\s*[_\s]", text_lower, re.MULTILINE))
        name_count = len(re.findall(r"(?:^|\s)(?:print\s+)?name:\s*[_\s]", text_lower, re.MULTILINE))
        if by_count + name_count >= 2:
            structural += 1

        if prev_is_candidate:
            structural += 1

    candidate.structural_score = structural
    candidate.total_score = primary_score + sig_block_score + structural

    # Qualification check
    qualifies = False
    if primary_score >= 4:  # 2+ primary marker hits
        qualifies = True
        candidate.detection_layer = "keyword-primary"
    elif primary_score >= 2 and structural >= 2:
        qualifies = True
        candidate.detection_layer = "keyword+structural"
    elif sig_block_score >= 3:
        qualifies = True
        candidate.detection_layer = "keyword-sigblock"
    elif is_notary:
        qualifies = True
        candidate.detection_layer = "notary"

    if not qualifies:
        candidate.total_score = 0  # Mark as not qualifying

    return candidate


def find_exhibit_start(page_texts: list[str]) -> int:
    """Find the first page index where exhibits begin.

    Returns the page index of the first exhibit heading, or len(page_texts) if none found.
    Uses multiple detection strategies since pandoc/PDF conversion may not preserve
    standalone exhibit headings from Word.
    """
    for i, text in enumerate(page_texts):
        lines = text.strip().split('\n')
        for line in lines[:8]:  # Check first several lines of each page
            stripped = line.strip()
            # Match standalone exhibit headings like "EXHIBIT A", "Exhibit B", "EXHIBITS"
            if re.match(r'^exhibit\s*[a-z0-9]?\s*$', stripped, re.IGNORECASE):
                return i
            # Match "EXHIBIT A — [Title]" style headings
            if re.match(r'^exhibit\s+[a-z0-9]\s*[-–—:]\s*', stripped, re.IGNORECASE):
                return i
            # Match "Exhibit A" appearing as a short line (heading-like)
            if re.match(r'^exhibit\s+[a-z0-9]\s*$', stripped, re.IGNORECASE) and len(stripped) < 20:
                return i
        # Also check if "Exhibit [X]" appears in flowing text at the start of a new section
        # (pandoc sometimes merges headings into text)
        text_lower = text.lower()
        if re.search(r'\bexhibit\s+[a-z]\b', text_lower[:200]) and i > 10:
            # Make sure it's not just a reference ("as defined in Exhibit A")
            # by checking if it's near the start of the page content
            first_200 = text_lower[:200]
            if re.match(r'\s*(?:\w+\s+){0,3}exhibit\s+[a-z]\b', first_200):
                return i
    return len(page_texts)


def is_exhibit_sig_page(page_text: str) -> bool:
    """Check if a page's footer indicates it's an exhibit signature page (to be excluded).

    Footer patterns like 'Signature Page - Exhibit A' or 'Signature - Exhibit B'
    signal that this is an exhibit sig page and should be skipped.
    """
    footer = get_footer_text(page_text).lower()
    return bool(re.search(r'signature?\s*(?:page)?\s*[-–—]\s*exhibit', footer, re.IGNORECASE))


def detect_sig_pages(pdf_path: Path) -> list[SigPageCandidate]:
    """Detect signature pages in a PDF file."""
    reader = PdfReader(str(pdf_path))
    total_pages = len(reader.pages)
    candidates = []

    # Extract all page texts
    page_texts = [extract_page_text(reader, i) for i in range(total_pages)]
    text_lengths = [len(t) for t in page_texts]
    avg_text_length = sum(text_lengths) / max(len(text_lengths), 1)

    # Find where exhibits start — sig pages in exhibits are excluded
    exhibit_start = find_exhibit_start(page_texts)
    if exhibit_start < total_pages:
        print(f"    Exhibits start at page {exhibit_start + 1} — excluding exhibit sig pages")

    # First pass: detect "[Signature Page Follows]" pages so we can boost the next page
    sig_page_follows = set()
    for i in range(total_pages):
        if re.search(r'\[?\s*signature\s+page\s+follows?\s*\]?', page_texts[i], re.IGNORECASE):
            sig_page_follows.add(i)

    prev_is_candidate = False
    for i in range(total_pages):
        # Skip pages in exhibit section (unless they have a non-exhibit footer override)
        if i >= exhibit_start:
            # Check if this page has an explicit "Signature - [Doc Name]" footer
            # (non-exhibit footer) that overrides exhibit exclusion
            footer_text = get_footer_text(page_texts[i])
            footer_match = check_footer_match(footer_text)
            if footer_match and not is_exhibit_sig_page(page_texts[i]):
                # Footer override — this sig page is explicitly marked even though
                # it's in the exhibit section
                pass
            else:
                # In exhibit territory and no override — skip
                if is_exhibit_sig_page(page_texts[i]):
                    print(f"    Skipping page {i + 1}: exhibit signature page (footer)")
                prev_is_candidate = False
                continue

        # Score the page
        candidate = score_page(page_texts[i], i, total_pages, avg_text_length, prev_is_candidate)
        candidate.source_file = str(pdf_path)

        # If the previous page had "[Signature Page Follows]", boost this page
        if (i - 1) in sig_page_follows and candidate.total_score > 0:
            candidate.structural_score += 3
            candidate.total_score += 3
            if "Follows prior [Sig Page Follows]" not in candidate.markers_found:
                candidate.markers_found.append("Follows prior [Sig Page Follows]")

        # If THIS page only has "[Signature Page Follows]" but no sig block markers,
        # it's a body page, not a sig page — demote it
        if i in sig_page_follows:
            has_sig_blocks = any(
                re.search(pattern, page_texts[i].lower(), re.MULTILINE)
                for pattern, _ in SIG_BLOCK_MARKERS
            )
            if not has_sig_blocks:
                candidate.total_score = 0  # Not a sig page itself

        # Handle combo sig+notary pages: if a page has both sig block markers AND
        # notary markers, classify it as SIG (not NOTARY) — the notary content is
        # part of the same page and will be included automatically
        if candidate.total_score > 0 and candidate.is_notary:
            has_primary = any(
                re.search(pattern, page_texts[i].lower())
                for pattern, _ in PRIMARY_MARKERS
            )
            has_sig_blocks = any(
                re.search(pattern, page_texts[i].lower(), re.MULTILINE)
                for pattern, _ in SIG_BLOCK_MARKERS
            )
            if has_primary or has_sig_blocks:
                # This is a combo page — classify as SIG, not NOTARY
                candidate.is_notary = False
                candidate.detection_layer = candidate.detection_layer or "keyword-combo"

        if candidate.total_score > 0:
            candidates.append(candidate)
            prev_is_candidate = True
        else:
            prev_is_candidate = False

    # Post-processing: gap detection for exhibit sig pages.
    # After the main sig block cluster, isolated sig-like pages separated by 3+ body
    # pages are likely exhibit sig pages and should be excluded (unless they have a
    # Layer 0 footer match, which is always trusted).
    if len(candidates) >= 2:
        filtered = []
        main_cluster_end = -1

        # Find the main sig cluster: consecutive or near-consecutive candidates
        # (within 2 pages of each other)
        cluster_start = candidates[0].page_number
        cluster_end = candidates[0].page_number
        for c in candidates[1:]:
            if c.page_number - cluster_end <= 2:
                cluster_end = c.page_number
            else:
                break
        main_cluster_end = cluster_end

        for c in candidates:
            if c.page_number <= main_cluster_end + 2:
                # In or near the main cluster — keep
                filtered.append(c)
            elif c.detection_layer == "footer":
                # Footer override — always trust
                filtered.append(c)
            elif c.page_number >= total_pages - 3:
                # Never exclude candidates on the last 3 pages — sig pages
                # are almost always at the end of the document
                filtered.append(c)
            else:
                # Isolated candidate far from main cluster — likely exhibit
                print(f"    Excluding page {c.page_number + 1}: "
                      f"likely exhibit (gap of {c.page_number - main_cluster_end} pages "
                      f"from main sig cluster)")

        candidates = filtered

    return candidates


# ---------------------------------------------------------------------------
# Phase 3: Signatory Name Extraction
# ---------------------------------------------------------------------------

def extract_signatory_names(page_text: str) -> list[str]:
    """Extract signatory names from a signature page."""
    names = []
    lines = page_text.split('\n')

    for i, line in enumerate(lines):
        stripped = line.strip()

        # Pattern 1: "Name:", "Print Name:", or "Printed:" followed by a name
        name_match = re.match(r'(?:print(?:ed)?\s+)?name:\s*(.+)', stripped, re.IGNORECASE)
        if not name_match:
            name_match = re.match(r'printed:\s*(.+)', stripped, re.IGNORECASE)
        if name_match:
            name_text = name_match.group(1).strip()
            # Skip if it's just underscores or blank
            if name_text and not re.match(r'^[_\s]+$', name_text):
                names.append(name_text)
                continue

        # Pattern 2: "personally appeared [NAME]" (notary pages)
        appeared_match = re.search(r'personally\s+appeared\s+(.+?)(?:\s*,|\s+known|\s+to\s+me)',
                                    stripped, re.IGNORECASE)
        if appeared_match:
            name_text = appeared_match.group(1).strip()
            if name_text and not re.match(r'^[_\s]+$', name_text):
                names.append(name_text)
                continue

        # Pattern 3: Entity name pattern "[ENTITY], a [state] [type]"
        entity_match = re.match(
            r'^(.+?),\s+a\s+(?:\w+\s+){0,3}(?:limited\s+liability\s+company|'
            r'corporation|partnership|limited\s+partnership|LLC|LP|Inc\.?|Corp\.?)',
            stripped, re.IGNORECASE
        )
        if entity_match:
            # Look for individual name in the "By:" line below
            for j in range(i + 1, min(i + 5, len(lines))):
                by_match = re.match(r'\s*by:\s*(.+)', lines[j].strip(), re.IGNORECASE)
                if by_match:
                    by_name = by_match.group(1).strip()
                    if by_name and not re.match(r'^[_\s]+$', by_name):
                        names.append(by_name)
                    break

    # Clean up names
    cleaned = []
    for name in names:
        # Remove trailing underscores, dates, common suffixes
        name = re.sub(r'[_]+.*$', '', name).strip()
        name = re.sub(r'\s*\(seal\)\s*$', '', name, flags=re.IGNORECASE).strip()
        # Skip if too short or looks like a title
        if len(name) < 2:
            continue
        if name.lower() in ("manager", "member", "president", "vice president",
                             "secretary", "treasurer", "director", "officer",
                             "authorized signatory", "authorized representative"):
            continue
        cleaned.append(name)

    return cleaned


def normalize_name(name: str) -> str:
    """Normalize a name to a canonical form for comparison."""
    # Strip extra whitespace, lowercase for comparison
    normalized = ' '.join(name.split()).strip()
    return normalized


def names_match(name1: str, name2: str, threshold: int = 80) -> bool:
    """Check if two names likely refer to the same person."""
    n1 = normalize_name(name1).lower()
    n2 = normalize_name(name2).lower()

    # Exact match
    if n1 == n2:
        return True

    # Fuzzy match
    score = fuzz.token_sort_ratio(n1, n2)
    return score >= threshold


def resolve_signatories(all_candidates: list[SigPageCandidate]) -> dict[str, SignatoryInfo]:
    """Resolve and normalize signatory names across all candidates."""
    signatories: dict[str, SignatoryInfo] = {}
    name_map: dict[str, str] = {}  # variant -> canonical

    for candidate in all_candidates:
        if candidate.is_notary:
            continue  # Notary pages get associated later

        names = extract_signatory_names(
            extract_page_text(PdfReader(candidate.source_file), candidate.page_number)
        )
        candidate.signatory_names = names

        for name in names:
            # Check if this name matches an existing signatory
            matched_canonical = None
            for variant, canonical in name_map.items():
                if names_match(name, variant):
                    matched_canonical = canonical
                    break

            if matched_canonical:
                signatories[matched_canonical].name_variants.append(name)
                signatories[matched_canonical].pages.append(candidate)
                name_map[name] = matched_canonical
            else:
                # New signatory — use the most complete version as canonical
                canonical = name
                signatories[canonical] = SignatoryInfo(
                    canonical_name=canonical,
                    name_variants=[name],
                    pages=[candidate]
                )
                name_map[name] = canonical

    return signatories


# ---------------------------------------------------------------------------
# Phase 4: Associate Notary Pages
# ---------------------------------------------------------------------------

def associate_notary_pages(
    candidates: list[SigPageCandidate],
    signatories: dict[str, SignatoryInfo]
) -> list[SigPageCandidate]:
    """Associate notary pages with their signatories."""
    unassociated = []

    for candidate in candidates:
        if not candidate.is_notary:
            continue

        page_text = extract_page_text(
            PdfReader(candidate.source_file), candidate.page_number
        )

        # Method 1: Name match from "personally appeared [NAME]"
        appeared_match = re.search(
            r'personally\s+appeared\s+(.+?)(?:\s*,|\s+known|\s+to\s+me)',
            page_text, re.IGNORECASE
        )
        matched = False
        if appeared_match:
            appeared_name = appeared_match.group(1).strip()
            for canonical, info in signatories.items():
                if any(names_match(appeared_name, v) for v in info.name_variants):
                    info.pages.append(candidate)
                    candidate.associated_signatory = canonical
                    matched = True
                    break

        # Method 2: Proximity — associate with preceding sig page in same document
        if not matched:
            same_doc_candidates = [
                c for c in candidates
                if c.source_file == candidate.source_file
                and not c.is_notary
                and c.page_number < candidate.page_number
            ]
            if same_doc_candidates:
                # Find the closest preceding sig page
                preceding = max(same_doc_candidates, key=lambda c: c.page_number)
                if preceding.signatory_names:
                    # Associate with the signatory of the preceding page
                    for name in preceding.signatory_names:
                        for canonical, info in signatories.items():
                            if any(names_match(name, v) for v in info.name_variants):
                                info.pages.append(candidate)
                                candidate.associated_signatory = canonical
                                matched = True
                                break
                        if matched:
                            break

        if not matched:
            unassociated.append(candidate)

    return unassociated


# ---------------------------------------------------------------------------
# Phase 5: Assemble PDFs
# ---------------------------------------------------------------------------

def sanitize_filename(name: str) -> str:
    """Sanitize a name for use in a filename."""
    # Remove special characters, replace spaces with hyphens
    sanitized = re.sub(r'[^\w\s-]', '', name)
    sanitized = re.sub(r'\s+', '-', sanitized).strip('-')
    return sanitized


def extract_pdf_page(source_pdf: str, page_idx: int, output_path: Path) -> bool:
    """Extract a single page from a PDF to a new PDF file."""
    try:
        reader = PdfReader(source_pdf)
        writer = PdfWriter()
        writer.add_page(reader.pages[page_idx])
        with open(output_path, 'wb') as f:
            writer.write(f)
        return True
    except Exception as e:
        print(f"  Error extracting page {page_idx} from {source_pdf}: {e}")
        return False


def build_bookmarked_pdf(
    pages: list[tuple[str, SigPageCandidate]],  # (doc_name, candidate)
    output_path: Path,
    temp_dir: Path
) -> bool:
    """Build a bookmarked PDF from a list of pages."""
    if not pages:
        return False

    writer = PdfWriter()
    current_doc = None
    doc_bookmark = None
    page_count = 0

    for doc_name, candidate in pages:
        reader = PdfReader(candidate.source_file)
        page = reader.pages[candidate.page_number]
        writer.add_page(page)

        if doc_name != current_doc:
            current_doc = doc_name
            doc_bookmark = writer.add_outline_item(doc_name, page_count)
        else:
            # Sub-bookmark for additional pages from same doc
            page_label = "Notary" if candidate.is_notary else f"Page {candidate.page_number + 1}"
            writer.add_outline_item(f"{doc_name} — {page_label}", page_count, parent=doc_bookmark)

        page_count += 1

    try:
        with open(output_path, 'wb') as f:
            writer.write(f)
        return True
    except Exception as e:
        print(f"  Error writing {output_path}: {e}")
        return False


# ---------------------------------------------------------------------------
# Main Pipeline
# ---------------------------------------------------------------------------

def run_pipeline(input_folder: str, auto_run: bool = False,
                 project_name_override: str = None) -> dict:
    """Run the full signature package extraction pipeline."""
    folder = Path(input_folder)
    if not folder.exists():
        print(f"Error: Folder not found: {input_folder}")
        sys.exit(1)

    results = {
        "project_name": "",
        "checklist_entries": [],
        "matched_files": [],
        "unmatched_checklist": [],
        "unmatched_files": [],
        "candidates_by_doc": {},
        "signatories": {},
        "files_created": [],
        "flags": [],
    }

    # --- Phase 1: Parse Checklist ---
    print("=" * 60)
    print("PHASE 1: Parsing Closing Checklist")
    print("=" * 60)

    # Project name: derive from filenames first (most reliable), then CLI override, then checklist
    filename_project_name = extract_project_name_from_filenames(folder)
    if project_name_override:
        project_name = project_name_override
    elif filename_project_name:
        project_name = filename_project_name
        print(f"Project name (from filenames): {project_name}")
    else:
        project_name = None  # Will try checklist next

    checklist_file = find_checklist_file(folder)
    if not checklist_file:
        print("No closing checklist found.")
        print("Please identify the checklist file, or type 'none' to process alphabetically.")
        results["flags"].append("No checklist found — processing alphabetically")
        # Fallback: alphabetical order
        all_files = sorted([
            f for f in folder.iterdir()
            if f.suffix.lower() in (".docx", ".pdf") and f.is_file()
        ])
        entries = [
            ChecklistEntry(order=i + 1, doc_name=f.stem, matched_file=str(f))
            for i, f in enumerate(all_files)
        ]
        if not project_name:
            project_name = "UnknownProject"
    else:
        print(f"Checklist found: {checklist_file.name}")
        checklist_project_name, entries = parse_checklist(checklist_file)
        if not project_name:
            project_name = checklist_project_name
        print(f"Project name: {project_name}")
        print(f"Documents listed: {len(entries)}")

        # Fuzzy match
        entries = fuzzy_match_files(entries, folder)

    results["project_name"] = project_name
    results["checklist_entries"] = entries

    # Report matching results
    print(f"\n--- Matching Results ---")
    for entry in entries:
        status = "✓" if entry.matched_file else "✗"
        matched_name = Path(entry.matched_file).name if entry.matched_file else "NO MATCH"
        score = f" ({entry.match_score}%)" if entry.matched_file else ""
        print(f"  {status} {entry.doc_name} → {matched_name}{score}")
        if entry.matched_file:
            results["matched_files"].append(entry)
        else:
            results["unmatched_checklist"].append(entry.doc_name)

    # Check for unmatched files (skip output files, temp files, and checklist)
    matched_paths = {e.matched_file for e in entries if e.matched_file}
    for f in folder.iterdir():
        if f.suffix.lower() in (".docx", ".pdf") and str(f) not in matched_paths:
            # Skip checklist, output files, temp/lock files
            if any(p in f.stem.lower() for p in CHECKLIST_PATTERNS):
                continue
            if "SigPackage" in f.name or "sig-package-results" in f.name:
                continue
            if f.name.startswith("~$"):
                continue
            results["unmatched_files"].append(f.name)
            print(f"  ? Unmatched file: {f.name}")

    # --- Phase 2: Detect Signature Pages ---
    print(f"\n{'=' * 60}")
    print("PHASE 2: Detecting Signature Pages")
    print("=" * 60)

    temp_dir = Path(tempfile.mkdtemp(prefix="sig-package-"))
    all_candidates = []

    for entry in entries:
        if not entry.matched_file:
            continue

        file_path = Path(entry.matched_file)
        print(f"\n  Scanning: {file_path.name}")

        # Convert .docx to PDF if needed
        if file_path.suffix.lower() == ".docx":
            pdf_path = convert_docx_to_pdf(file_path, temp_dir)
            if not pdf_path:
                results["flags"].append(f"Failed to convert {file_path.name} to PDF")
                continue
        else:
            pdf_path = file_path

        candidates = detect_sig_pages(pdf_path)

        # Store the original source info
        for c in candidates:
            c.source_file = str(pdf_path)

        results["candidates_by_doc"][entry.doc_name] = candidates
        all_candidates.extend(candidates)

        if candidates:
            print(f"    Found {len(candidates)} sig page candidate(s):")
            for c in candidates:
                page_type = "NOTARY" if c.is_notary else "SIG"
                print(f"      [{page_type}] Page {c.page_number + 1}: "
                      f"markers={c.markers_found}, score={c.total_score}")
        else:
            print(f"    No sig pages detected")
            results["flags"].append(f"No sig pages: {file_path.name}")

    if not auto_run:
        # Review mode: output summary for user review
        print(f"\n--- REVIEW MODE: Candidate Summary ---")
        print(f"{'Doc Name':<40} {'Page':<6} {'Type':<8} {'Layer':<20} {'Score':<6} {'Markers'}")
        print("-" * 120)
        for entry in entries:
            if entry.doc_name in results["candidates_by_doc"]:
                for c in results["candidates_by_doc"][entry.doc_name]:
                    page_type = "NOTARY" if c.is_notary else "SIG"
                    markers_str = ", ".join(c.markers_found[:3])
                    if len(c.markers_found) > 3:
                        markers_str += f" (+{len(c.markers_found) - 3} more)"
                    print(f"  {entry.doc_name[:38]:<40} {c.page_number + 1:<6} "
                          f"{page_type:<8} {c.detection_layer:<20} {c.total_score:<6.0f} {markers_str}")
        print("\nReview the candidates above. Confirm to proceed with assembly.")
        print("(In auto-run mode, this step is skipped.)")

    # --- Phase 3: Extract Signatory Names ---
    print(f"\n{'=' * 60}")
    print("PHASE 3: Extracting Signatory Names")
    print("=" * 60)

    signatories = resolve_signatories(all_candidates)
    results["signatories"] = signatories

    for canonical, info in signatories.items():
        print(f"  Signatory: {canonical}")
        print(f"    Variants: {info.name_variants}")
        print(f"    Pages: {len(info.pages)}")

    # Check for unsigned pages
    unsigned_candidates = [
        c for c in all_candidates
        if not c.is_notary and not c.signatory_names
    ]
    if unsigned_candidates:
        print(f"\n  Unsigned pages: {len(unsigned_candidates)}")

    # --- Phase 4: Associate Notary Pages ---
    print(f"\n{'=' * 60}")
    print("PHASE 4: Associating Notary Pages")
    print("=" * 60)

    unassociated_notary = associate_notary_pages(all_candidates, signatories)
    if unassociated_notary:
        results["flags"].append(
            f"{len(unassociated_notary)} notary page(s) could not be associated with a signatory"
        )
        for c in unassociated_notary:
            print(f"  Unassociated notary: {Path(c.source_file).name} page {c.page_number + 1}")

    # --- Phase 5: Assemble PDFs ---
    print(f"\n{'=' * 60}")
    print("PHASE 5: Assembling PDF Packages")
    print("=" * 60)

    project_slug = sanitize_filename(project_name)

    # Build document order map for sorting
    doc_order = {}
    for entry in entries:
        if entry.matched_file:
            doc_order[entry.matched_file] = entry.order
            # Also map the PDF conversion path
            pdf_name = Path(entry.matched_file).stem + ".pdf"
            pdf_path = temp_dir / pdf_name
            if pdf_path.exists():
                doc_order[str(pdf_path)] = entry.order

    def sort_key(item):
        """Sort by checklist order."""
        _, candidate = item
        return doc_order.get(candidate.source_file, 999)

    # Find doc name for a candidate
    def get_doc_name(candidate):
        for entry in entries:
            if entry.matched_file:
                matched_pdf = temp_dir / (Path(entry.matched_file).stem + ".pdf")
                if str(matched_pdf) == candidate.source_file or entry.matched_file == candidate.source_file:
                    return entry.doc_name
        return Path(candidate.source_file).stem

    # 1. Combined package
    combined_pages = []
    for candidate in all_candidates:
        doc_name = get_doc_name(candidate)
        combined_pages.append((doc_name, candidate))
    combined_pages.sort(key=sort_key)

    combined_path = folder / f"{project_slug}-Combined-SigPackage.pdf"
    if build_bookmarked_pdf(combined_pages, combined_path, temp_dir):
        results["files_created"].append(str(combined_path))
        print(f"  Created: {combined_path.name} ({len(combined_pages)} pages)")

    # 2. Per-signatory packages
    for canonical, info in signatories.items():
        sig_pages = []
        for candidate in info.pages:
            doc_name = get_doc_name(candidate)
            sig_pages.append((doc_name, candidate))
        sig_pages.sort(key=sort_key)

        sig_slug = sanitize_filename(canonical)
        sig_path = folder / f"{project_slug}-{sig_slug}-SigPackage.pdf"
        if build_bookmarked_pdf(sig_pages, sig_path, temp_dir):
            results["files_created"].append(str(sig_path))
            print(f"  Created: {sig_path.name} ({len(sig_pages)} pages)")

    # 3. Unsigned package
    if unsigned_candidates:
        unsigned_pages = []
        for candidate in unsigned_candidates:
            doc_name = get_doc_name(candidate)
            unsigned_pages.append((doc_name, candidate))
        unsigned_pages.sort(key=sort_key)

        unsigned_path = folder / f"{project_slug}-Unsigned-SigPackage.pdf"
        if build_bookmarked_pdf(unsigned_pages, unsigned_path, temp_dir):
            results["files_created"].append(str(unsigned_path))
            print(f"  Created: {unsigned_path.name} ({len(unsigned_pages)} pages)")

    # Include unassociated notary pages in combined only (already there)
    # Flag them
    for c in unassociated_notary:
        results["flags"].append(
            f"Unassociated notary page included in combined package only: "
            f"{Path(c.source_file).name} page {c.page_number + 1}"
        )

    # --- Report ---
    print(f"\n{'=' * 60}")
    print("RESULTS SUMMARY")
    print("=" * 60)
    print(f"Project: {project_name}")
    print(f"Documents processed: {len([e for e in entries if e.matched_file])}")
    print(f"Sig pages found: {len(all_candidates)}")
    print(f"Signatories identified: {len(signatories)}")
    print(f"Unsigned pages: {len(unsigned_candidates)}")
    print(f"\nFiles created:")
    for f in results["files_created"]:
        print(f"  {Path(f).name}")
    if results["flags"]:
        print(f"\nFlags:")
        for flag in results["flags"]:
            print(f"  ⚠ {flag}")

    # Write results JSON
    results_json_path = folder / f"{project_slug}-sig-package-results.json"
    json_results = {
        "project_name": project_name,
        "documents_processed": len([e for e in entries if e.matched_file]),
        "sig_pages_found": len(all_candidates),
        "signatories": {k: {"name": v.canonical_name, "page_count": len(v.pages)}
                        for k, v in signatories.items()},
        "unsigned_pages": len(unsigned_candidates),
        "files_created": [Path(f).name for f in results["files_created"]],
        "flags": results["flags"],
        "unmatched_checklist": results["unmatched_checklist"],
        "unmatched_files": results["unmatched_files"],
    }
    with open(results_json_path, 'w') as f:
        json.dump(json_results, f, indent=2)
    print(f"\nResults saved to: {results_json_path.name}")

    return results


# ---------------------------------------------------------------------------
# CLI Entry Point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Extract signature pages from closing documents and build PDF packages."
    )
    parser.add_argument("input_folder", help="Path to folder containing closing documents")
    parser.add_argument("--auto-run", action="store_true",
                        help="Skip review mode and auto-assemble")
    parser.add_argument("--project-name", type=str, default=None,
                        help="Override the project name (otherwise extracted from checklist)")

    args = parser.parse_args()
    run_pipeline(args.input_folder, auto_run=args.auto_run,
                 project_name_override=args.project_name)


if __name__ == "__main__":
    main()
